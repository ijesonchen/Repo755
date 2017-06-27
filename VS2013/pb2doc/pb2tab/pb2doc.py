from docx import Document


s = r'''


package ULPB;

//////////////////////////////////////////////////////////////////////////
// common message
message RetCode
{
    required    int32   code = 1;    // 返回值。0：成功，非0：出现错误
    optional    string  msg = 2;    // 出现错误时的详细信息
}

//////////////////////////////////////////////////////////////////////////
// call process

// 11种: vad, front-end (vadseg), dtmf, lang, gender, spk, sseg, kwd, trans, de-ring, de-duplicate
enum AudioRecogFlag
{
    ARF_Unknown                         = 0; // 未定义
    ARF_Preprocess                      = 1;  // 有效音判定，语种，性别
    ARF_Speaker                         = 2;  // 说话人
    ARF_Keyword                         = 4;  // 关键词
    ARF_Transcript                      = 8;  // 全文
    ARF_VAD                             = 16; // 端点检测（有效音片段）
    ARF_SpeechSegment                   = 32; // 固定音频
    ARF_DTMF                            = 64; // 拨号音
    ARF_DeRing                          = 512;  // 振铃
    ARF_RtTrans                         = 1024; // 实时转录
}

enum GenderType
{
    GT_Male     = 1; // 男
    GT_Female   = 2; // 女
    GT_Both     = 3; // 男和女
}

enum ChannelType
{
    CT_Unknown     = 0; // 未定义
    CT_Pcm         = 1; // PCM
    CT_Voip        = 2; // VoIP
    CT_TwoLines    = 3; // 两线
    CT_AirSignal   = 4; // 空口
    CT_WX          = 5; // 微信
    CT_Mico        = 41; // 微波
    CT_Satelite    = 42; // 卫星
    CT_VHF         = 43;   // 超高频
    CT_HF          = 44;   // 高频
    CT_FIBER       = 45;   // 光线
}

enum AudioFormat
{
    AT_Unknown         = 0; // 未定义
    AT_ALaw            = 1; // A率
    AT_Gsm610          = 2; // GSM610
    AT_G726            = 3; // G726
    AT_LinearPCM       = 4; // 线性PCM
    AT_G723_16         = 5; // G723.1
    AT_G723_15         = 6; // G723.2
    AT_G729            = 7; // G729
    AT_ULaw            = 8; // u率
    AT_INTEL_G723      = 9; // INTEL_G723
    AT_INTEL_G729      = 10; // INTEL_G729
    AT_INTEL_GSM690    = 11; // INTEL_GSM690
    AT_AMR_NB          = 12; // AMR_NB
    AT_GSM_EFR         = 13; // GSM_EFR
    AT_STEREO_PCM      = 41;   // STEREO_PCM
    AT_STEREO_ALaw     = 42;   // STEREO_ALaw
}

enum FileTag
{
    AFT_Unknown     = 0; // 未定义
    AFT_Combination = 1; // 合路
    AFT_ClusterA    = 2; // A路
    AFT_ClusterB    = 3; // B路
    AFT_ClusterC    = 4; // C路
}

//////////////////////////////////////////////////////////////////////////
// search
enum SearchType
{
    ST_CondToVoice                     = 1; // 条件检索语音
    ST_VoiceToVoice                    = 2; // 语音检索语音
    ST_VoiceToSpeaker                  = 3; // 语音检索对象
}

enum SearchState
{
    SS_Unknown                        = 0; // 未定义
    SS_Queue                          = 1; // 队列中
    SS_Running                        = 2; // 正在运行
    SS_Finish                         = 3; // 完成
    SS_Abort                          = 4; // 取消
    SS_Error                          = 5; // 错误
}

message SearchTaskInfo
{
    optional    SearchState   runstate  = 1; // 任务状态
    optional    uint64  queuetime       = 2; // 排队时间
    optional    uint64  starttime       = 3; // 开始时间
    optional    uint64  finishtime      = 4; // 结束时间
    optional    uint64  aborttime       = 5; // 取消时间
    optional    RetCode  ret            = 6; // 操作结果描述
}


message PhoneNumber
{
    optional    string  telnumber       = 1;    // 电话号码
    optional    string  countrycode     = 2;    // 国家码
    optional    string  provincecode    = 3;    // 省份码
    optional    string  areacode        = 4;    // 地区码
}

// 结果类型
enum CallResultType
{
    CallResultNorm      = 0x0000;   // 正常结果
    CallResultExclude   = 0x0001;   // 排除结果
    CallResultFraud     = 0x0002;   // 诈骗结果
}


//////////////////////////////////////////////////////////////////////////
// file result items

// search: v2v
message FileItem
{
    optional    string  callid  = 1; // 业务系统话单id
    optional    string  spyid   = 2;    // 业务系统id
    optional    FileTag filetag = 3; // 归属标志（合成/A/B/C路语音）
    optional    float   score   = 4; // 输出值
}

message PreItem
{
    optional    bool    isValid     = 1; // 是否有效音
    optional    string  lang        = 2; // 语种
    optional    float   langscore   = 3;    // 语种输出
    optional    GenderType  gender  = 4; // 性别
    optional    uint32  validSecs   = 5;    // 有效音长
    optional    uint32  callSecs    = 6;    // 通话音长
}

message SpkItem
{
    optional    string  objectid        = 1; // 对象id
    optional    float   score           = 2; // 输出
    optional    string  spyid           = 3; // 上线资源的系统id
}

message VadItem
{
    optional    int32   begin           = 1; // 开始时间，毫秒数
    optional    int32   end             = 2; // 结束时间，毫秒数
    optional    string  spk             = 3; // 说话人标识：有A，B，C三个选项,配合合路语音分离使用
}

message DtmfItem
{
    optional    int32   code            = 1; // 拨号音代号
    optional    int32   begin           = 2; // 开始时间，毫秒数
    optional    int32   end             = 3; // 结束时间，毫秒数
}

message SSegItem
{
    required    string  ssegid      = 1; // 固定音频id
    required    int32   ssegbegin   = 2; // 固定音频中的开始时间，毫秒数
    required    int32   ssegend     = 3; // 固定音频中的结束时间，毫秒数
    required    int32   filebegin   = 4; // 文件中的开始时间，毫秒数
    required    int32   fileend     = 5; // 文件中的结束时间，毫秒数
    required    float   score       = 6; // 输出值
    optional    string  spyid       = 7; // 设置资源的系统id
}

message KwdItem
{
    optional    string  keyword     = 1; // 关键词文本
    optional    string  lang        = 2; // 语种
    optional    float   score       = 3; // 输出值
    optional    uint32  begintime   = 4; // 开始时间，毫秒数
    optional    uint32  endtime     = 5; // 结束时间，毫秒数
    optional    string  text        = 6; // 关键词所在全文文本
    optional    string  spyid       = 7; // 设置资源的系统id
}

message TransItem
{
    optional    string  text    = 1; // 文本
    optional    float   score   = 3; // 输出值
    optional    uint32  begintime   = 4; // 开始时间，毫秒数
    optional    uint32  endtime     = 5; // 结束时间，毫秒数
}

message DupItem
{
    optional    string  callid   = 1; // 业务系统话单id
    optional    float   score    = 2; // 输出值
}

message RingItem
{
    optional int32 begin  = 1; // 开始时间，毫秒数
    optional int32 end    = 2; // 结束时间，毫秒数
}

// tag: vad lang gender
// part: spk
// fast: front-end, dtmf, sseg, 
// slow: kwd, trans
// other: de-ring, de-duplicate
// for real-time result
message RTResult
{
    optional    string  callid  = 1; // 业务系统话单id
    optional    FileTag filetag = 2; // 归属标志（合成/A/B/C路语音）
    optional    string  spyid   = 3; // 业务系统id
    optional    RetCode status  = 4; // 操作结果描述
    // partial result
    optional    PreItem prelist  = 5;    // 单个文件的结果
    repeated    SpkItem speakers = 6; // 实时预警说话人结果
    // fast result
    repeated    VadItem     vads   = 7; // 有效音片段信息
    repeated    DtmfItem    dtmfs  = 8; // 拨号音信息
    repeated    SSegItem    ssegs  = 9; // 实时预警固定音频结果
    // slow result
    repeated    KwdItem     kwds    = 10; // 实时预警关键词结果
    repeated    TransItem   trans   = 11; // 全文转写结果
    // other
    repeated    DupItem     dups    = 12; // 去重结果
    repeated    RingItem    rings   = 13; // 振铃检测结果

    // 语音识别类型标志
    // 即使处理过单没结果仍然置相应位，通过判断结果是否为空确定是否有相应结果
    optional    uint32  resultFlag          = 14;  // 识别类型标志
}

// for history search result
message HistResult
{
    optional    string  callid  = 1; // 业务系统话单id
    optional    FileTag filetag = 2; // 归属标志（合成/A/B/C路语音）
    optional    string  spyid   = 3; //  业务系统id 
    repeated    SpkItem spks    = 4; // 说话人结果
    repeated    SSegItem ssegs  = 5; // 固定音频结果
    repeated    KwdItem  kwds   = 6; // 关键词结果
}

//////////////////////////////////////////////////////////////////////////
// resource management

enum OnlineCate
{
    OnlineCateNormal    = 0; // 正常对象
    OnlineCateExclude   = 1; // 排除对象
    OnlineCateFraud     = 2; // 诈骗对象
}

enum OperationType
{
    OT_Unknown          = 0; // 未定义
    OT_Create           = 1; // 新建
    OT_Delete           = 2; // 删除
    OT_Update           = 3; // 更新
}


message TrainInfo
{
    optional    AudioFormat code    = 1; // 文件编码格式
    optional    bool        isdata  = 2; // 是否传输数据
    optional    bytes       voice   = 3; // 语音数据
    optional    string      path    = 4; // 语音路径
    optional    float       score   = 5; // 评估样本输出
}

enum RuleType
{
    RT_Unknown          = 0; // 未定义
    RT_Speaker          = 1; // 说话人
    RT_Keyword          = 2; // 关键词
    RT_SSeg             = 3; // 固定音频
}


message SpkRule
{
    optional    string  objectid        = 1; // 说话人id
    optional    string  spyid           = 2; // 业务系统id
    optional    float   threshold       = 3; // 阈值
    repeated    PhoneNumber telrule     = 4; // 号码策略
    optional    OnlineCate  category    = 5; // 上线类型
}

message SSegRule
{
    optional    string  speechsegid     = 1; // 固定音频id
    optional    bytes   data            = 2; // 固定音频数据
    optional    string  speechsegname   = 3; // 名称
    optional    float   threshold       = 4; // 阈值
    optional    string  lang            = 5; // 语种
    optional    GenderType gender       = 6; // 性别
    repeated    PhoneNumber telrule     = 7; // 号码策略
    optional    OnlineCate  category    = 8; // 上线类型
}

message KeywordRule
{
    optional    string  lang            = 1; // 语种
    optional    string  keyword         = 2; // 关键词
    optional    float   threshold       = 3; // 阈值
    repeated    PhoneNumber telrule     = 5; // 号码策略
}




//////////////////////////////////////////////////////////////////////////
// alarm & status
enum AlarmLevel
{
    AL_Info         = 0; // 状态信息
    AL_General      = 1; // 一般告警
    AL_Important    = 2; // 重要告警
    AL_Urgent       = 3; // 紧急告警
}

message StateItem
{
    required    string  name    = 1; // 字段名
    required    string  value   = 2; // 字段值
    optional    string  remark  = 3; // 备注
}

enum RecogLineTag
{
    RecogLineDefault = 0; // 默认（ABC路都识别）
    RecogLineA = 1; // 仅识别A路
    RecogLineB = 2; // 仅识别B路
    RecogLineB = 4; // 仅识别C路
}


//////////////////////////////////////////////////////////////////////////
// pretreat
enum MachineType {
    VprPreMT = 11; // vad,lang,gender,vpr
    VadPreMT = 12; // vadseg,dtmf,sseg
    KwdPreMT = 13; // kwd,trans,mand
    WewrPreMT = 15; // kwd,trans,wewr
    DuplicateMT = 16;  // dup
    HuyuPreMT = 17; // kwd,trans,huyu
}

enum OnlineTag
{
    OnlineTagKwd = 1; // 正常关键词
    OnlineTagSpk = 2; // 正常说话人
    OnlineTagSeg = 4; // 正常固定音频
    OnlineTagSpkExclude = 8; // 排除说话人
    OnlineTagSegExclude = 16; // 排除固定音频
    OnlineTagSpkFraud = 32; // 诈骗说话人
    OnlineTagSegFraud = 64; // 诈骗固定音频
}

message SpeakerInfo
{
    optional string guid    = 1; // 说话人guid
    optional string spyid   = 2; // 上线资源的业务系统id
    optional string objectid = 3; // 业务系统中说话人id
    optional string modelpath = 4; // 识别系统中模型路径
    optional bool   isAutoTrain = 5; // 是否自动训练对象

    optional OnlineCate category = 6;   // 上线类型
    optional uint64 timestamp = 7; // 最后更新时间
    // filter
    optional float  threshold = 8; // 阈值
    optional GenderType gender = 9; // 性别
    optional string lang = 10;   // 语种
    repeated PhoneNumber telrule = 11; // 号码策略
    repeated string dataSources = 12; // 要匹配的业务系统
}

message SSegInfo
{
    optional string guid    = 1; // 固定音频guid
    optional string spyid   = 2;   // 上线资源的业务系统id
    optional string ssegid  = 3; // 业务系统中固定音频id
    optional string vocpath = 4; // 语音识别系统中固定音频路径

    optional OnlineCate category = 5;   // 上线类型
    optional uint64 timestamp = 6; // 最后更新时间

    // filter
    optional float  threshold = 7; // 阈值
    optional string lang      = 8;   // 可选
    optional GenderType gender = 9; // 性别
    repeated PhoneNumber telrule = 11; // 号码策略
    repeated string dataSources = 12;   //  要匹配的业务系统
}

message KeywordInfo
{
    optional string guid = 1; // 关键词guid
    optional string spyid = 2; // 上线资源的业务系统id
    optional string lang = 3; // 语种
    optional string keyword = 4; // 关键词文本
    optional uint64 timestamp = 5; // 最后更新时间

    // filter
    optional float  threshold = 6; // 阈值
    repeated PhoneNumber telrule = 8; // 号码策略
    repeated string dataSources = 9 ; // 要匹配的业务系统
}


'''

A1 =  1
A2 = 1000000
A3 = 1000000
A4 =  1


document = Document()

# 1 msg 2 enum
T = 0

table = None

a = s.split('\n')
for l in a:
    l = l.strip()
    if l[:2] == '//':
        continue
    s2 = l.replace(';', ' ').replace('//', ' ').replace('=', ' ')
    a2 = s2.split()
    if len(a2) == 2:
        if a2[0] == 'message':
            T = 1
            document.add_paragraph()
            # title
            s = ('%s %s'%(a2[0],a2[1]))
            #print(s)
            p = document.add_heading(a2[1], level=3)
            TBL_COL = 5
            table = document.add_table(rows=1, cols=TBL_COL, style='Table Grid')
            hdr_cells = table.rows[0].cells
            for i in range(TBL_COL - 1):
                hdr_cells[0].merge(hdr_cells[i + 1])
            hdr_cells[0].text = '消息类型： ' + s
            row_cells = table.add_row().cells
            row_cells[0].text = '限定符'
            row_cells[0].width = A1
            row_cells[1].text = '类型'
            row_cells[1].width = A2
            row_cells[2].text = '名字'
            row_cells[2].width = A3
            row_cells[3].text = 'ID'
            row_cells[3].width = A4
            row_cells[4].text = '说明'
        elif a2[0] == 'enum':
            T = 2
            document.add_paragraph()
            # title
            s = ('%s %s'%(a2[0],a2[1]))
            #print(s)
            p = document.add_heading(a2[1], level=3)
            TBL_COL = 3
            table = document.add_table(rows=1, cols=TBL_COL, style='Table Grid')
            hdr_cells = table.rows[0].cells
            for i in range(TBL_COL - 1):
                hdr_cells[0].merge(hdr_cells[i + 1])
            hdr_cells[0].text = '枚举类型： ' + s
            row_cells = table.add_row().cells
            row_cells[0].text = '名字'
            row_cells[0].width = A1
            row_cells[1].text = 'ID'
            row_cells[1].width = A4
            row_cells[2].text = '说明'
        else:
            T = 0
            s = ('%s %s'%(a2[0],a2[1]))
            #print(s)
            print('!!!error', s)
    if len(a2) == 5:
        if T != 1:
            print('error content')
            continue
        #for s3 in a2:
        #    print(s3, end=' ')
        #print()
        
        row_cells = table.add_row().cells
        row_cells[0].text = a2[0]
        row_cells[0].width = A1
        row_cells[1].text = a2[1]
        row_cells[1].width = A2
        row_cells[2].text = a2[2]
        row_cells[2].width = A3
        row_cells[3].text = a2[3]
        row_cells[3].width = A4
        row_cells[4].text = a2[4]
        
    if len(a2) == 3:
        if T != 2:
            print('error content')
            continue
        #for s3 in a2:
        #    print(s3, end=' ')
        #print()
        
        row_cells = table.add_row().cells
        row_cells[0].text = a2[0]
        row_cells[0].width = A1
        row_cells[1].text = a2[1]
        row_cells[1].width = A4
        row_cells[2].text = a2[2]


document.add_paragraph()
document.save('demo2.docx')


