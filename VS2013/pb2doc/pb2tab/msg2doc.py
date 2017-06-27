from docx import Document

document = Document()




p = document.add_paragraph('A plain paragraph having some ')
table = document.add_table(rows=1, cols=3, style='Table Grid')
hdr_cells = table.rows[0].cells
x = hdr_cells[0]
y = hdr_cells[1]
z = hdr_cells[2]
x.merge(y)
x.merge(z)
hdr_cells[0].text = 'Qty'
for i in range(5):
    row_cells = table.add_row().cells
    row_cells[0].text = 'qty' + str(i)
    row_cells[1].text = 'id' + str(i)
    row_cells[2].text = 'desc' + str(i)

p = document.add_paragraph()
p = document.add_heading('title3', level=3)

table = document.add_table(rows=1, cols=3, style='Table Grid')
hdr_cells = table.rows[0].cells
for i in range(2):
    hdr_cells[0].merge(hdr_cells[i + 1])
hdr_cells[0].text = 'Qty'
for i in range(5):
    row_cells = table.add_row().cells
    row_cells[0].text = 'qty' + str(i)
    row_cells[1].text = 'id' + str(i)
    row_cells[2].text = 'desc' + str(i)


document.save('demo.docx')