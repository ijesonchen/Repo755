from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
x = hdr_cells[0]
y = hdr_cells[1]
z = hdr_cells[2]
x.merge(y)
x.merge(z)
hdr_cells[0].text = 'Qty'
for i in range(5):
    row_cells = table.add_row().cells
    row_cells[0].text = 'qty' + str(i)
    row_cells[1].text = 'id' + str(i)
    row_cells[2].text = 'desc' + str(i)

document.add_page_break()

document.save('demo.docx')